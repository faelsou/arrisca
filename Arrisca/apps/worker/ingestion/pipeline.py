"""Pipeline de ingestão de documentos.

Fluxo (executado pela task Celery `ingestion.ingest_document`):

    1. Marca o documento como `processing`.
    2. Baixa o arquivo do `source_uri` (file:// ou http(s)://).
    3. Extrai texto (pypdf para PDF, python-docx para DOCX, leitura
       direta para TXT/MD), preservando o número da página quando
       disponível (melhora as citações no chat).
    4. Chunka por tokens (~800 tokens, overlap de 100), respeitando
       fronteiras de parágrafo sempre que possível.
    5. Gera embeddings em lote (OpenAI text-embedding-3-small).
    6. Persiste em transação única: remove chunks antigos do documento
       (caso de re-ingestão) e insere os novos com os metadados
       DENORMALIZADOS do documento pai (tenant_id, area_id,
       sensitivity, is_current) — ver comentário no schema.
    7. Marca o documento como `completed` (ou `failed`, com o erro).

Decisões:
    * asyncpg com conexão dedicada por task (não o pool da API — o
      worker roda em processo próprio e ingestão é de baixa frequência).
    * O embedding vai como literal string `[...]` com cast `::vector`
      no SQL — dispensa registrar codec do pgvector na conexão.
    * tiktoken (cl100k_base) para contagem de tokens; se indisponível,
      fallback de ~4 chars/token.
"""

from __future__ import annotations

import logging
import os
import re
from dataclasses import dataclass, field
from pathlib import Path
from urllib.parse import urlparse

import asyncpg
import httpx
from openai import AsyncOpenAI

log = logging.getLogger("arrisca.ingestion")

# ---------------------------------------------------------------------------
# Configuração (via env, com defaults sensatos)
# ---------------------------------------------------------------------------

EMBEDDING_MODEL = os.getenv("EMBEDDING_MODEL", "text-embedding-3-small")
EMBEDDING_DIM = int(os.getenv("EMBEDDING_DIM", "1536"))
CHUNK_TOKENS = int(os.getenv("CHUNK_TOKENS", "800"))
CHUNK_OVERLAP_TOKENS = int(os.getenv("CHUNK_OVERLAP_TOKENS", "100"))
EMBEDDING_BATCH_SIZE = int(os.getenv("EMBEDDING_BATCH_SIZE", "100"))
DOWNLOAD_TIMEOUT_S = float(os.getenv("INGESTION_DOWNLOAD_TIMEOUT", "60"))
MAX_DOWNLOAD_BYTES = int(os.getenv("INGESTION_MAX_BYTES", str(50 * 1024 * 1024)))  # 50 MB

# ---------------------------------------------------------------------------
# Tokenizer (tiktoken com fallback)
# ---------------------------------------------------------------------------

try:  # pragma: no cover - depende do ambiente
    import tiktoken

    _ENCODER = tiktoken.get_encoding("cl100k_base")

    def count_tokens(text: str) -> int:
        return len(_ENCODER.encode(text))

except Exception:  # noqa: BLE001 - fallback deliberado
    _ENCODER = None

    def count_tokens(text: str) -> int:
        # Aproximação conservadora: ~4 caracteres por token em pt-BR/en.
        return max(1, len(text) // 4)


# ---------------------------------------------------------------------------
# Tipos internos
# ---------------------------------------------------------------------------


@dataclass(slots=True)
class ExtractedBlock:
    """Bloco de texto extraído do documento, com página de origem (1-based)."""

    text: str
    page: int | None = None


@dataclass(slots=True)
class Chunk:
    """Chunk pronto para embedding + insert."""

    index: int
    content: str
    token_count: int
    metadata: dict = field(default_factory=dict)


class IngestionError(RuntimeError):
    """Erro de ingestão com mensagem amigável (vai para documents.ingestion_error)."""


# ---------------------------------------------------------------------------
# 1. Download
# ---------------------------------------------------------------------------


async def download_source(source_uri: str) -> bytes:
    """Baixa o conteúdo bruto do documento a partir do source_uri.

    Suporta:
        file://  — caminho local montado no container do worker
        http(s):// — URL pública ou pré-assinada (S3 presigned funciona aqui)
    """
    if not source_uri:
        raise IngestionError("Documento sem source_uri — nada para ingerir.")

    parsed = urlparse(source_uri)

    if parsed.scheme == "file":
        path = Path(parsed.path)
        if not path.is_file():
            raise IngestionError(f"Arquivo não encontrado no worker: {path}")
        data = path.read_bytes()

    elif parsed.scheme in ("http", "https"):
        async with httpx.AsyncClient(
            timeout=DOWNLOAD_TIMEOUT_S, follow_redirects=True
        ) as client:
            resp = await client.get(source_uri)
            if resp.status_code != 200:
                raise IngestionError(
                    f"Download falhou com HTTP {resp.status_code} para {source_uri}"
                )
            data = resp.content

    elif parsed.scheme == "s3":
        raise IngestionError(
            "source_uri s3:// direto não é suportado — gere uma URL pré-assinada "
            "(https://) no upload, ou adicione boto3 e implemente aqui."
        )
    else:
        raise IngestionError(f"Esquema de source_uri não suportado: {parsed.scheme!r}")

    if len(data) > MAX_DOWNLOAD_BYTES:
        raise IngestionError(
            f"Arquivo excede o limite de {MAX_DOWNLOAD_BYTES // (1024 * 1024)} MB."
        )
    if not data:
        raise IngestionError("Download retornou arquivo vazio.")
    return data


# ---------------------------------------------------------------------------
# 2. Extração de texto
# ---------------------------------------------------------------------------

_EXT_BY_CONTENT_TYPE = {
    "application/pdf": ".pdf",
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document": ".docx",
    "text/plain": ".txt",
    "text/markdown": ".md",
}


def _infer_kind(source_uri: str, content_type: str | None) -> str:
    """Retorna 'pdf' | 'docx' | 'text' a partir do content_type ou da extensão."""
    ext = ""
    if content_type:
        ext = _EXT_BY_CONTENT_TYPE.get(content_type.split(";")[0].strip().lower(), "")
    if not ext and source_uri:
        ext = Path(urlparse(source_uri).path).suffix.lower()

    if ext == ".pdf":
        return "pdf"
    if ext == ".docx":
        return "docx"
    if ext in (".txt", ".md", ".markdown", ".csv", ".log", ""):
        return "text"
    raise IngestionError(
        f"Formato não suportado (content_type={content_type!r}, extensão={ext!r}). "
        "Suportados: PDF, DOCX, TXT, MD."
    )


def extract_blocks(data: bytes, kind: str) -> list[ExtractedBlock]:
    """Extrai blocos de texto do binário conforme o tipo do documento."""
    if kind == "pdf":
        return _extract_pdf(data)
    if kind == "docx":
        return _extract_docx(data)
    return _extract_text(data)


def _extract_pdf(data: bytes) -> list[ExtractedBlock]:
    import io

    from pypdf import PdfReader

    try:
        reader = PdfReader(io.BytesIO(data))
    except Exception as exc:  # noqa: BLE001
        raise IngestionError(f"PDF inválido ou corrompido: {exc}") from exc

    if reader.is_encrypted:
        try:
            reader.decrypt("")  # tenta senha vazia
        except Exception as exc:  # noqa: BLE001
            raise IngestionError("PDF protegido por senha — remova a proteção.") from exc

    blocks: list[ExtractedBlock] = []
    for page_num, page in enumerate(reader.pages, start=1):
        try:
            text = page.extract_text() or ""
        except Exception:  # noqa: BLE001 - página problemática não derruba o doc
            log.warning("Falha ao extrair página %d — pulando.", page_num)
            continue
        text = _normalize_whitespace(text)
        if text:
            blocks.append(ExtractedBlock(text=text, page=page_num))
    return blocks


def _extract_docx(data: bytes) -> list[ExtractedBlock]:
    import io

    from docx import Document as DocxDocument

    try:
        doc = DocxDocument(io.BytesIO(data))
    except Exception as exc:  # noqa: BLE001
        raise IngestionError(f"DOCX inválido ou corrompido: {exc}") from exc

    parts: list[str] = []
    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            parts.append(text)
    # Tabelas: cada linha vira "cel1 | cel2 | cel3"
    for table in doc.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells]
            line = " | ".join(c for c in cells if c)
            if line:
                parts.append(line)

    full = _normalize_whitespace("\n\n".join(parts))
    return [ExtractedBlock(text=full)] if full else []


def _extract_text(data: bytes) -> list[ExtractedBlock]:
    for encoding in ("utf-8", "latin-1"):
        try:
            text = data.decode(encoding)
            break
        except UnicodeDecodeError:
            continue
    else:  # pragma: no cover
        raise IngestionError("Não foi possível decodificar o arquivo de texto.")
    text = _normalize_whitespace(text)
    return [ExtractedBlock(text=text)] if text else []


def _normalize_whitespace(text: str) -> str:
    text = text.replace("\r\n", "\n").replace("\r", "\n")
    text = re.sub(r"[ \t]+", " ", text)
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text.strip()


# ---------------------------------------------------------------------------
# 3. Chunking (por tokens, respeitando parágrafos)
# ---------------------------------------------------------------------------


def chunk_blocks(
    blocks: list[ExtractedBlock],
    *,
    max_tokens: int = CHUNK_TOKENS,
    overlap_tokens: int = CHUNK_OVERLAP_TOKENS,
) -> list[Chunk]:
    """Divide os blocos em chunks de ~max_tokens com overlap.

    Estratégia: acumula parágrafos até estourar o limite; ao fechar um
    chunk, o próximo começa com a cauda (overlap) do anterior. Um
    parágrafo maior que o limite é fatiado por sentenças/limite bruto.
    """
    # (parágrafo, página) achatados na ordem do documento
    paragraphs: list[tuple[str, int | None]] = []
    for block in blocks:
        for para in block.text.split("\n\n"):
            para = para.strip()
            if para:
                paragraphs.append((para, block.page))

    chunks: list[Chunk] = []
    current: list[tuple[str, int | None]] = []
    current_tokens = 0

    def flush() -> None:
        nonlocal current, current_tokens
        if not current:
            return
        content = "\n\n".join(p for p, _ in current)
        pages = [pg for _, pg in current if pg is not None]
        meta: dict = {}
        if pages:
            meta["page_start"] = min(pages)
            meta["page_end"] = max(pages)
        chunks.append(
            Chunk(
                index=len(chunks),
                content=content,
                token_count=count_tokens(content),
                metadata=meta,
            )
        )
        # Prepara o overlap: mantém parágrafos do fim até ~overlap_tokens
        tail: list[tuple[str, int | None]] = []
        tail_tokens = 0
        for para, pg in reversed(current):
            t = count_tokens(para)
            if tail_tokens + t > overlap_tokens:
                break
            tail.insert(0, (para, pg))
            tail_tokens += t
        # Se nenhum parágrafo inteiro coube (parágrafos > overlap_tokens),
        # cai para as últimas sentenças do último parágrafo.
        if not tail and current:
            last_para, last_pg = current[-1]
            sentences = re.split(r"(?<=[.!?])\s+", last_para)
            sent_tail: list[str] = []
            for sent in reversed(sentences):
                t = count_tokens(sent)
                if tail_tokens + t > overlap_tokens:
                    break
                sent_tail.insert(0, sent)
                tail_tokens += t
            if sent_tail:
                tail = [(" ".join(sent_tail), last_pg)]
        current = tail
        current_tokens = tail_tokens

    for para, page in paragraphs:
        para_tokens = count_tokens(para)

        # Parágrafo gigante: fatia por sentenças
        if para_tokens > max_tokens:
            flush()
            for piece in _split_long_paragraph(para, max_tokens):
                current.append((piece, page))
                current_tokens += count_tokens(piece)
                if current_tokens >= max_tokens:
                    flush()
            continue

        if current_tokens + para_tokens > max_tokens:
            flush()
        current.append((para, page))
        current_tokens += para_tokens

    flush()
    return chunks


def _split_long_paragraph(text: str, max_tokens: int) -> list[str]:
    """Fatia um parágrafo maior que o limite por sentenças; se uma
    sentença sozinha estourar, corta por janela bruta de caracteres."""
    sentences = re.split(r"(?<=[.!?])\s+", text)
    pieces: list[str] = []
    buf = ""
    for sent in sentences:
        candidate = f"{buf} {sent}".strip() if buf else sent
        if count_tokens(candidate) > max_tokens and buf:
            pieces.append(buf)
            buf = sent
        else:
            buf = candidate
        # Sentença sozinha ainda gigante → corte bruto
        while count_tokens(buf) > max_tokens:
            approx_chars = max_tokens * 4
            pieces.append(buf[:approx_chars])
            buf = buf[approx_chars:]
    if buf:
        pieces.append(buf)
    return pieces


# ---------------------------------------------------------------------------
# 4. Embeddings (OpenAI, em lote)
# ---------------------------------------------------------------------------


async def embed_chunks(chunks: list[Chunk]) -> list[list[float]]:
    """Gera embeddings em lotes. Ordem do retorno = ordem dos chunks."""
    client = AsyncOpenAI()  # OPENAI_API_KEY via env; retries automáticos do SDK
    embeddings: list[list[float]] = []

    for start in range(0, len(chunks), EMBEDDING_BATCH_SIZE):
        batch = chunks[start : start + EMBEDDING_BATCH_SIZE]
        resp = await client.embeddings.create(
            model=EMBEDDING_MODEL,
            input=[c.content for c in batch],
        )
        # A API preserva a ordem, mas ordenamos por index por segurança.
        for item in sorted(resp.data, key=lambda d: d.index):
            embeddings.append(item.embedding)
        log.info(
            "Embeddings: %d/%d chunks processados.",
            min(start + EMBEDDING_BATCH_SIZE, len(chunks)),
            len(chunks),
        )

    if len(embeddings) != len(chunks):
        raise IngestionError(
            f"OpenAI retornou {len(embeddings)} embeddings para {len(chunks)} chunks."
        )
    for emb in embeddings:
        if len(emb) != EMBEDDING_DIM:
            raise IngestionError(
                f"Dimensão inesperada de embedding: {len(emb)} (esperado {EMBEDDING_DIM})."
            )
    return embeddings


def _embedding_to_pgvector(embedding: list[float]) -> str:
    return "[" + ",".join(f"{x:.7f}" for x in embedding) + "]"


# ---------------------------------------------------------------------------
# 5. Persistência
# ---------------------------------------------------------------------------


async def _connect() -> asyncpg.Connection:
    dsn = os.getenv("DATABASE_URL")
    if not dsn:
        raise IngestionError("DATABASE_URL não configurada no worker.")
    return await asyncpg.connect(dsn=dsn)


async def _mark_processing(conn: asyncpg.Connection, document_id: str) -> asyncpg.Record:
    row = await conn.fetchrow(
        """
        UPDATE documents
           SET ingestion_status = 'processing',
               ingestion_error  = NULL,
               updated_at       = NOW()
         WHERE id = $1
        RETURNING id, tenant_id, area_id, sensitivity, is_current,
                  source_uri, content_type, title
        """,
        document_id,
    )
    if row is None:
        raise IngestionError(f"Documento {document_id} não existe.")
    return row


async def _persist_chunks(
    conn: asyncpg.Connection,
    doc: asyncpg.Record,
    chunks: list[Chunk],
    embeddings: list[list[float]],
) -> None:
    """Transação única: apaga chunks antigos e insere os novos."""
    import json

    async with conn.transaction():
        await conn.execute(
            "DELETE FROM document_chunks WHERE document_id = $1", doc["id"]
        )
        await conn.executemany(
            """
            INSERT INTO document_chunks (
                document_id, tenant_id, area_id, sensitivity, is_current,
                chunk_index, content, embedding, token_count, metadata
            ) VALUES ($1, $2, $3, $4, $5, $6, $7, $8::vector, $9, $10::jsonb)
            """,
            [
                (
                    doc["id"],
                    doc["tenant_id"],
                    doc["area_id"],
                    doc["sensitivity"],
                    doc["is_current"],
                    chunk.index,
                    chunk.content,
                    _embedding_to_pgvector(emb),
                    chunk.token_count,
                    json.dumps(chunk.metadata),
                )
                for chunk, emb in zip(chunks, embeddings)
            ],
        )
        await conn.execute(
            """
            UPDATE documents
               SET ingestion_status = 'completed',
                   ingested_at      = NOW(),
                   updated_at       = NOW()
             WHERE id = $1
            """,
            doc["id"],
        )


async def _mark_failed(document_id: str, error: str) -> None:
    """Melhor esforço para registrar a falha (conexão própria, fora da transação)."""
    try:
        conn = await _connect()
        try:
            await conn.execute(
                """
                UPDATE documents
                   SET ingestion_status = 'failed',
                       ingestion_error  = $2,
                       updated_at       = NOW()
                 WHERE id = $1
                """,
                document_id,
                error[:2000],
            )
        finally:
            await conn.close()
    except Exception:  # noqa: BLE001 - não mascarar o erro original
        log.exception("Não foi possível marcar o documento %s como failed.", document_id)


# ---------------------------------------------------------------------------
# Entrypoint
# ---------------------------------------------------------------------------


async def run_ingestion(document_id: str) -> dict:
    """Executa o pipeline completo para um documento. Retorna um resumo."""
    conn = await _connect()
    try:
        doc = await _mark_processing(conn, document_id)

        log.info("[%s] baixando %s", document_id, doc["source_uri"])
        data = await download_source(doc["source_uri"])

        kind = _infer_kind(doc["source_uri"] or "", doc["content_type"])
        log.info("[%s] extraindo texto (%s, %d bytes)", document_id, kind, len(data))
        blocks = extract_blocks(data, kind)
        if not blocks:
            raise IngestionError(
                "Nenhum texto extraído — o documento pode ser escaneado (imagem). "
                "OCR não é suportado neste pipeline."
            )

        chunks = chunk_blocks(blocks)
        if not chunks:
            raise IngestionError("Chunking não produziu nenhum chunk.")
        log.info("[%s] %d chunks gerados", document_id, len(chunks))

        embeddings = await embed_chunks(chunks)

        await _persist_chunks(conn, doc, chunks, embeddings)
        log.info("[%s] ingestão concluída: %d chunks", document_id, len(chunks))
        return {
            "document_id": document_id,
            "chunks": len(chunks),
            "tokens": sum(c.token_count for c in chunks),
        }
    except Exception as exc:
        await _mark_failed(document_id, str(exc))
        raise
    finally:
        await conn.close()


__all__ = [
    "run_ingestion",
    "download_source",
    "extract_blocks",
    "chunk_blocks",
    "embed_chunks",
    "IngestionError",
]
